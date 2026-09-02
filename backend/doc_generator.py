# 25010 Characteristic: Reliability

import os
import sys
import platform
import subprocess
import re
import copy
from docx import Document
from docx.shared import Pt
from docx.table import _Row
from datetime import datetime
from database import BASE_DIR

def log_conversion_error(msg):
    try:
        log_path = os.path.join(BASE_DIR, "PASADA_CRASH_LOG.txt")
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(f"[{datetime.now()}] DOC_GEN ERROR: {msg}\n")
    except:
        pass

def get_resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(os.path.dirname(__file__))
    return os.path.join(base_path, relative_path)

def replace_text_in_paragraph(paragraph, mapping):
    full_text = paragraph.text
    found_keys = [k for k in mapping.keys() if k in full_text]
    
    if not found_keys:
        return
        
    base_font_name = paragraph.runs[0].font.name if paragraph.runs else None
    base_font_size = paragraph.runs[0].font.size if paragraph.runs else None

    for run in paragraph.runs:
        run.text = ""

    pattern = re.compile("|".join(map(re.escape, found_keys)))
    parts = pattern.split(full_text)
    matches = pattern.findall(full_text)
    
    for i in range(len(parts)):
        if parts[i]:
            run = paragraph.add_run(parts[i])
            if base_font_name: run.font.name = base_font_name
            if base_font_size: run.font.size = base_font_size
            
        if i < len(matches):
            key = matches[i]
            val = str(mapping[key])
            run = paragraph.add_run(val)
            run.bold = True 
            
            if key in ["[SBN_NO]", "{{SBN_NO}}"]:
                run.font.size = Pt(12)
            elif base_font_size: 
                run.font.size = base_font_size
                
            if base_font_name: run.font.name = base_font_name

def clean_val(v):
    if not v: return ""
    s = str(v).strip().upper()
    if s in ["NAN", "NONE", "N/A", "UNKNOWN", "NULL"]: return ""
    return s

def generate_certificate(data: dict, settings: dict, template_path: str = "template.docx", output_dir: str = "exports", return_format: str = "pdf"):
    out_path = os.path.join(BASE_DIR, output_dir)
    if not os.path.exists(out_path):
        os.makedirs(out_path, exist_ok=True)

    actual_template_path = get_resource_path(template_path)
    
    try:
        doc = Document(actual_template_path)
    except Exception as e:
        raise FileNotFoundError(f"Template not found at {actual_template_path}. Error: {e}")

    issue_date_obj = data.get("issue_date") or datetime.now()
    valid_until_obj = data.get("valid_until") or datetime(issue_date_obj.year, 12, 31)

    replacements = {
        "[SBN_NO]": data.get("sbn_no", ""),
        "{{SBN_NO}}": data.get("sbn_no", ""),
        "[NAME]": data.get("operator_name", "").upper(),
        "{{NAME}}": data.get("operator_name", "").upper(),
        "[ADDRESS]": data.get("address", "").upper(),
        "{{ADDRESS}}": data.get("address", "").upper(),
        "[MOTOR_NO]": clean_val(data.get("motor_no")),
        "{{MOTOR_NO}}": clean_val(data.get("motor_no")),
        "[CHASSIS_NO]": clean_val(data.get("chassis_no")),
        "{{CHASSIS_NO}}": clean_val(data.get("chassis_no")),
        "[MAKE]": clean_val(data.get("make")),
        "{{MAKE}}": clean_val(data.get("make")),
        "[PLATE_NO]": clean_val(data.get("plate_no")),
        "{{PLATE_NO}}": clean_val(data.get("plate_no")),
        "[ROUTE]": data.get("driving_route", "").upper(),
        "{{ROUTE}}": data.get("driving_route", "").upper(),
        "[ISSUE_DATE]": issue_date_obj.strftime("%B %d, %Y"),
        "{{ISSUE_DATE}}": issue_date_obj.strftime("%B %d, %Y"),
        "[VALID_UNTIL]": valid_until_obj.strftime("%B %d, %Y"),
        "{{VALID_UNTIL}}": valid_until_obj.strftime("%B %d, %Y"),
        "[CHAIRMAN_NAME]": settings.get("committee_chair", "RODRIGO A. CASTILLO").upper(),
        "{{CHAIRMAN_NAME}}": settings.get("committee_chair", "RODRIGO A. CASTILLO").upper()
    }

    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

    raw_name = data.get('operator_name', 'Unknown')
    safe_name = re.sub(r'[^A-Za-z0-9_\-]', '', str(raw_name).replace(' ', '_'))
    
    docx_path = os.path.abspath(os.path.join(out_path, f"MTOP_{safe_name}.docx"))
    doc.save(docx_path)

    if return_format == "docx":
        return docx_path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    pdf_path = os.path.abspath(os.path.join(out_path, f"MTOP_{safe_name}.pdf"))
    
    try:
        if platform.system() == "Windows":
            try:
                import pythoncom
                import win32com.client
                pythoncom.CoInitialize() 
                word = win32com.client.DispatchEx("Word.Application")
                try:
                    word.Visible = False
                    word.DisplayAlerts = 0
                    doc_obj = word.Documents.Open(docx_path, ReadOnly=True)
                    doc_obj.SaveAs(pdf_path, FileFormat=17)
                    doc_obj.Close()
                finally:
                    word.Quit()
                    pythoncom.CoUninitialize()
                if os.path.exists(pdf_path): return pdf_path, "application/pdf"
            except Exception as e:
                log_conversion_error(f"Windows COM PDF conversion failed: {e}. Trying LibreOffice...")

        subprocess.run(
            ['libreoffice', '--headless', '--nologo', '--nofirststartwizard', '--convert-to', 'pdf', docx_path, '--outdir', out_path],
            check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
        )
        if os.path.exists(pdf_path): 
            return pdf_path, "application/pdf"
            
    except subprocess.CalledProcessError as e:
        log_conversion_error(f"LibreOffice PDF conversion failed: {e}. Falling back to .docx.")
    except FileNotFoundError as e:
        log_conversion_error(f"LibreOffice not found: {e}. Falling back to .docx.")
    except Exception as e:
        log_conversion_error(f"Unexpected PDF conversion error: {e}. Falling back to .docx.")
        
    return docx_path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def generate_toda_summary(summary: dict, settings: dict, template_path: str = "dashboard_template.docx", output_dir: str = "exports", return_format: str = "docx"):
    out_path = os.path.join(BASE_DIR, output_dir)
    if not os.path.exists(out_path):
        os.makedirs(out_path, exist_ok=True)

    actual_template_path = get_resource_path(template_path)
    
    try:
        doc = Document(actual_template_path)
    except Exception as e:
        raise FileNotFoundError(f"Template not found at {actual_template_path}. Error: {e}")

    as_of_date = str(summary.get("as_of_date", datetime.now().strftime("%B %d, %Y")))
    year = summary.get("year", datetime.now().year)
    
    header_replacements = {
        "[AS_OF_DATE]": as_of_date,
        "{{AS_OF_DATE}}": as_of_date,
        "[TOTAL_TODA]": str(summary.get("total_toda", 0)),
        "{{TOTAL_TODA}}": str(summary.get("total_toda", 0)),
        "[GRAND_TOTAL]": str(summary.get("grand_total", 0)),
        "{{GRAND_TOTAL}}": str(summary.get("grand_total", 0)),
        "[GRAND_RENEWED]": str(summary.get("grand_renewed", 0)),
        "{{GRAND_RENEWED}}": str(summary.get("grand_renewed", 0)),
        "[CHAIRMAN_NAME]": settings.get("committee_chair", "RODRIGO A. CASTILLO").upper(),
        "{{CHAIRMAN_NAME}}": settings.get("committee_chair", "RODRIGO A. CASTILLO").upper()
    }

    # Replace headers in paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, header_replacements)

    # Find the row containing the [ROUTE_NAME] token
    marker_row = None
    marker_table = None
    for table in doc.tables:
        for row in table.rows:
            row_text = "".join(c.text for c in row.cells)
            if "[ROUTE_NAME]" in row_text or "{{ROUTE_NAME}}" in row_text:
                marker_row = row
                marker_table = table
                break
        if marker_row:
            break

    # Clone the row and insert the data
    if marker_table and marker_row:
        for i, r_data in enumerate(summary.get("rows", []), start=1):
            new_tr = copy.deepcopy(marker_row._tr)
            # FIX: Insert BEFORE the marker row so it stays above the "TOTAL" row
            marker_row._tr.addprevious(new_tr)
            new_row = _Row(new_tr, marker_table)
            
            row_reps = {
                "[ROW_NO]": str(i),
                "{{ROW_NO}}": str(i),
                "[ROUTE_NAME]": str(r_data.get("route", "")),
                "{{ROUTE_NAME}}": str(r_data.get("route", "")),
                "[ROUTE_TOTAL]": str(r_data.get("total", 0)),
                "{{ROUTE_TOTAL}}": str(r_data.get("total", 0)),
                "[ROUTE_RENEWED]": str(r_data.get("renewed", 0)),
                "{{ROUTE_RENEWED}}": str(r_data.get("renewed", 0))
            }
            
            for cell in new_row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, row_reps)

        # Delete the original placeholder row
        marker_table._tbl.remove(marker_row._tr)
        
    # Final pass to catch [GRAND_TOTAL] and other header replacements inside the table
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, header_replacements)

    base_filename = f"TOTAL RENEWAL {year}"
    docx_path = os.path.abspath(os.path.join(out_path, f"{base_filename}.docx"))
    doc.save(docx_path)

    if return_format == "docx":
        return docx_path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    pdf_path = os.path.abspath(os.path.join(out_path, f"{base_filename}.pdf"))
    
    # PDF Conversion fallback matching MTOP
    try:
        if platform.system() == "Windows":
            try:
                import pythoncom
                import win32com.client
                pythoncom.CoInitialize() 
                word = win32com.client.DispatchEx("Word.Application")
                try:
                    word.Visible = False
                    word.DisplayAlerts = 0
                    doc_obj = word.Documents.Open(docx_path, ReadOnly=True)
                    doc_obj.SaveAs(pdf_path, FileFormat=17)
                    doc_obj.Close()
                finally:
                    word.Quit()
                    pythoncom.CoUninitialize()
                if os.path.exists(pdf_path): return pdf_path, "application/pdf"
            except Exception as e:
                log_conversion_error(f"Windows COM PDF conversion failed: {e}. Trying LibreOffice...")

        subprocess.run(
            ['libreoffice', '--headless', '--nologo', '--nofirststartwizard', '--convert-to', 'pdf', docx_path, '--outdir', out_path],
            check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
        )
        if os.path.exists(pdf_path): 
            return pdf_path, "application/pdf"
            
    except subprocess.CalledProcessError as e:
        log_conversion_error(f"LibreOffice PDF conversion failed: {e}. Falling back to .docx.")
    except FileNotFoundError as e:
        log_conversion_error(f"LibreOffice not found: {e}. Falling back to .docx.")
    except Exception as e:
        log_conversion_error(f"Unexpected PDF conversion error: {e}. Falling back to .docx.")
        
    return docx_path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"