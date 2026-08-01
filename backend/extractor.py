import re
import io
from docx import Document
import pandas as pd

def clean_field(text):
    if not text: return ""
    # Violently strips any overlapping labels that got caught by sloppy table formatting
    labels_to_strip = [r"Plate\s*No\.?[\:\s]*", r"Route[\:\s]*", r"Chassis\s*No\.?[\:\s]*", r"Make[\:\s]*", r"TERMS.*"]
    for label in labels_to_strip:
        text = re.sub(rf"(?i){label}.*$", "", text)
    return text.strip()

def extract_docx_data(contents, default_route, current_year):
    doc = Document(io.BytesIO(contents))
    
    text_blocks = []
    for p in doc.paragraphs:
        if p.text.strip(): text_blocks.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip(): text_blocks.append(p.text.strip())
                    
    # Inject pipe walls to stop fields from merging
    full_text = " || ".join(text_blocks)
    full_text = re.sub(r'\s+', ' ', full_text)
    
    def get_match(pattern):
        match = re.search(pattern, full_text, re.IGNORECASE)
        return match.group(1).strip() if match else ""

    # Strict lookaheads to isolate exact values
    op_name = get_match(r"NAME\s*[\:\.]+\s*(.*?)(?=\|\||ADDRESS|Motor)")
    address = get_match(r"ADDRESS\s*[\:\.]+\s*(.*?)(?=\|\||Motor|Chassis|Make)")
    raw_motor = get_match(r"Motor\s*No\.?[\s\:]+(.*?)(?=\|\||Plate|Chassis|Make)")
    raw_plate = get_match(r"Plate\s*No\.?[\s\:]+(.*?)(?=\|\||Chassis|Route|Make)")
    raw_chassis = get_match(r"Chassis\s*No\.?[\s\:]+(.*?)(?=\|\||Route|Make)")
    make = get_match(r"Make\s*[\:\.]+\s*(.*?)(?=\|\||TERMS)")
    physical_route = get_match(r"Route\s*[\:\.]+\s*(.*?)(?=\|\||Make|Plate|TERMS)")

    # Run through the scrubber to prevent duplicate labels
    raw_motor = clean_field(raw_motor)
    raw_plate = clean_field(raw_plate)
    raw_chassis = clean_field(raw_chassis)
    make = clean_field(make)
    physical_route = clean_field(physical_route)

    # UPDATED: Fault-Tolerant Universal SBN Extractor (Ignores spaces, handles optional years)
    sbn_match = re.search(r"([A-Z]{2,5}\s*[\-\–]\s*\d{3,}(?:\s*[\-\–]\s*\d{2,4})?)", full_text, re.IGNORECASE)
    sbn_no = re.sub(r'\s+', '', sbn_match.group(1)).replace('–', '-').strip() if sbn_match else f"{default_route[:3]}-000-{str(current_year)[-2:]}"

    date_match = re.search(r"(?:Date Issued|Given this)[:\s]*([a-zA-Z]+\s+\d{1,2},?\s+\d{4})", full_text, re.IGNORECASE)
    try:
        issue_date = pd.to_datetime(date_match.group(1)).to_pydatetime() if date_match else None
    except:
        issue_date = None

    return {
        "sbn_no": sbn_no.upper(),
        "operator_name": op_name.upper(),
        "address": address.upper(),
        "motor_no": raw_motor.upper(),
        "plate_no": raw_plate.upper(),
        "chassis_no": raw_chassis.upper(),
        "make": make.upper(),
        "driving_route": physical_route.upper() if physical_route else "POBLACION",
        "issue_date": issue_date
    }